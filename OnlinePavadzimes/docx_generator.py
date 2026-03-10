from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import os

def fmt_curr(val):
    return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", " ")

def add_horizontal_line(doc):
    """Izveido horizontālu līniju Word dokumentā, izmantojot krāsotu 1x1 tabulu."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17)
    
    # Iestatām fona krāsu
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="CDBF96"/>'.format(nsdecls('w')))
    tcPr.append(shd)
    
    # Iestatām ļoti mazu augstumu
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = parse_xml(r'<w:trHeight {} w:val="20" w:hRule="exact"/>'.format(nsdecls('w')))
    trPr.append(trHeight)
    
    doc.add_paragraph() # Atstarpe pēc līnijas

def generate_docx(data):
    doc = Document()
    
    # --- Dokumenta apmales ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    # --- Pamatstils (Arial, lai izskatītos līdzīgi PDF Roboto) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # ==========================================
    # 1. LOGO UN DOKUMENTA INFO
    # ==========================================
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)
    
    cell_logo = table.cell(0, 0)
    paragraph = cell_logo.paragraphs[0]
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_filename = "BRATUS MELNS LOGO PNG.png"
    logo_path = os.path.join(current_dir, logo_filename)

    try:
        paragraph.add_run().add_picture(logo_path, width=Cm(3.5))
    except Exception as e:
        paragraph.add_run("LOGO").bold = True
        
    cell_info = table.cell(0, 1)
    p = cell_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc_type = data.get('doc_type', 'Pavadzīme')
    doc_id = data.get('doc_id', 'BR 0000')
    date = data.get('date', '')
    due_date = data.get('due_date', '')
    
    # Virsraksts 14pt un Treknrakstā
    run = p.add_run(f"{doc_type} Nr. {doc_id}\n")
    run.bold = True
    run.font.size = Pt(14)
    
    p.add_run(f"\nDatums: {date}\n")
    p.add_run(f"Apmaksāt līdz: {due_date}")
    
    doc.add_paragraph()
    add_horizontal_line(doc)
    
    # ==========================================
    # 2. KLIENTS
    # ==========================================
    p = doc.add_paragraph()
    p.add_run("KLIENTS").bold = True
    
    p = doc.add_paragraph()
    p.add_run(data.get('client_name', '')).bold = True
    p.add_run(f"\nAdrese: {data.get('client_address', '')}").italic = True
    p.add_run(f"\nReģ. Nr.: {data.get('client_reg_no', '')}").italic = True
    p.add_run(f"\nPVN Nr.: {data.get('client_vat_no', '')}").italic = True
    
    doc.add_paragraph()
    add_horizontal_line(doc)
    
    # ==========================================
    # 3. PIEGĀDĀTĀJS UN BANKA
    # ==========================================
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)
    
    # Sender
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.add_run("SIA Bratus").bold = True
    p.add_run("\nAdrese: Ķekavas nov., Ķekava,").italic = True
    p.add_run("\nDārzenieku iela 42, LV-2123").italic = True
    p.add_run("\nReģ. Nr.: 40203628316").italic = True
    p.add_run("\nPVN Nr.: LV40203628316").italic = True
    p.add_run("\nTālrunis: +371 24424434").italic = True
    
    # Bank
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    run_bank = p.add_run("AS Swedbank")
    run_bank.bold = True
    run_bank.italic = True
    p.add_run("\nSWIFT/BIC: HABALV22").italic = True
    p.add_run("\nBankas konta numurs: ").italic = True
    p.add_run("LV64HABA0551060367591").bold = True
    
    doc.add_paragraph()
    
    # ==========================================
    # 4. PREČU TABULA
    # ==========================================
    headers = ["NOSAUKUMS", "Mērvienība", "DAUDZUMS", "CENA (EUR)", "KOPĀ (EUR)"]
    items = data.get('items', [])
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    table.columns[0].width = Cm(6.5)
    table.columns[1].width = Cm(2.5)
    table.columns[2].width = Cm(2.5)
    table.columns[3].width = Cm(2.5)
    table.columns[4].width = Cm(3.0)
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        tcPr = hdr_cells[i]._element.tcPr
        shd = parse_xml(r'<w:shd {} w:fill="CDBF96"/>'.format(nsdecls('w')))
        tcPr.append(shd)
        p = hdr_cells[i].paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item['name']
        row_cells[1].text = item['unit']
        row_cells[2].text = str(item['qty'])
        row_cells[3].text = str(item['price'])
        row_cells[4].text = str(item['total'])
        
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_paragraph()
    
    # ==========================================
    # 5. KOPSUMMAS
    # ==========================================
    subtotal = data.get('subtotal', '0.00')
    vat = data.get('vat', '0.00')
    total = data.get('total', '0.00')
    raw_discount_eur = data.get('raw_discount_eur', 0.0)
    
    is_advance_doc = ("avansa" in doc_type.lower())
    
    num_rows = 5 if raw_discount_eur > 0 else 3

    # Veidojam tabulu BEZ apmalēm (Table Normal stils to nodrošina automātiski)
    table = doc.add_table(rows=num_rows, cols=3)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(5.5)
    table.columns[2].width = Cm(3)

    def set_total_row(row_idx, label, value, label_bold, val_bold):
        cell_lbl = table.cell(row_idx, 1)
        cell_val = table.cell(row_idx, 2)
        p = cell_lbl.paragraphs[0]
        p.add_run(label).bold = label_bold
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        p = cell_val.paragraphs[0]
        p.add_run(f"{value} €").bold = val_bold
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if is_advance_doc:
        if raw_discount_eur > 0:
            set_total_row(0, "KOPĀ (bez PVN un atlaides)", subtotal, False, False)
            set_total_row(1, f"Atlaides apjoms ({data.get('discount_percent', 0):g}%)", f"-{data.get('discount_eur', '0.00')}", False, False)
            set_total_row(2, "Kopā ar atlaidi (bez PVN)", data.get('subtotal_after_discount', '0.00'), False, False)
            set_total_row(3, "PVN", vat, False, False)
            set_total_row(4, "KOPUMĀ APMAKSAI", total, False, False)
        else:
            set_total_row(0, "KOPĀ", subtotal, False, False)
            set_total_row(1, "PVN", vat, False, False)
            set_total_row(2, "Kopumā", total, False, False)
    else:
        if raw_discount_eur > 0:
            set_total_row(0, "KOPĀ (bez PVN un atlaides)", subtotal, True, False)
            set_total_row(1, f"Atlaides apjoms ({data.get('discount_percent', 0):g}%)", f"-{data.get('discount_eur', '0.00')}", True, False)
            set_total_row(2, "Kopā ar atlaidi (bez PVN)", data.get('subtotal_after_discount', '0.00'), True, False)
            set_total_row(3, "PVN", vat, True, False)
            set_total_row(4, "KOPUMĀ APMAKSAI", total, True, True)
        else:
            set_total_row(0, "KOPĀ", subtotal, True, False)
            set_total_row(1, "PVN", vat, True, False)
            set_total_row(2, "Kopumā", total, True, True)
    
    # ==========================================
    # 6. SUMMA VĀRDIEM UN AVANSS
    # ==========================================
    if is_advance_doc:
        doc.add_paragraph()
        raw_advance = data.get('raw_advance', 0.0)
        formatted_advance = fmt_curr(raw_advance)
        percent_val = int(round(data.get('advance_percent', 0)))
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"APMAKSĀJAMAIS AVANSS ({percent_val}%): {formatted_advance} €").bold = True

    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = "Vārdiem: "
    p.add_run(f"{prefix}{data.get('amount_words', '')}").italic = True
    
    doc.add_paragraph()
    add_horizontal_line(doc)
    
    # ==========================================
    # 7. PARAKSTI UN PAPILDINFO
    # ==========================================
    p_info = doc.add_paragraph()
    p_info.add_run("Papildus informācija:").bold = True
    
    comments = data.get('comments', '').strip()
    if comments:
        p_info.add_run(f"\n{comments}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    signatory = data.get('signatory', 'SIA Bratus valdes loceklis Adrians Stankevičs')
    
    if "pavadzīme" in doc_type.lower():
        prepared_text = f"Pavadzīmi sagatavoja: "
        received_text = "Pavadzīmi saņēma:"
    elif "avansa rēķins" in doc_type.lower():
        prepared_text = f"Avansa rēķinu sagatavoja: "
        received_text = "Avansa rēķinu saņēma:"
    else:
        prepared_text = f"Rēķinu sagatavoja: "
        received_text = "Rēķinu saņēma:"
        
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(7)
    
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.add_run(prepared_text)
    p.add_run(f"{signatory}").italic = True
    
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.add_run("__________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    p.add_run(received_text)
    
    cell = table.cell(1, 1)
    p = cell.paragraphs[0]
    p.add_run("__________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
